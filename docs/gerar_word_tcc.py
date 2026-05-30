#!/usr/bin/env python3
"""
gerar_word_tcc.py
=================
Converte TCC_BEDFLOW_ATLAS_COMPLETO.txt em documento .docx formatado conforme
o padrão PUCRS identificado no PDF de referência.

Uso:
    python3 docs/gerar_word_tcc.py
"""

import re
import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
except ImportError:
    print("Instalando python-docx...")
    os.system("pip install python-docx -q")
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

# ─── caminhos ─────────────────────────────────────────────────────────────────
BASE = os.path.dirname(os.path.abspath(__file__))
TXT_IN   = os.path.join(BASE, "TCC_BEDFLOW_ATLAS_COMPLETO.txt")
DOCX_OUT = os.path.join(BASE, "TCC_BEDFLOW_ATLAS_COMPLETO.docx")

# ─── cores e fontes ───────────────────────────────────────────────────────────
BLUE  = RGBColor(0x00, 0x46, 0x7F)   # azul PUCRS (headings)
BLACK = RGBColor(0x00, 0x00, 0x00)
GRAY  = RGBColor(0x55, 0x55, 0x55)   # notas / legendas
FONT  = "Calibri"
MONO  = "Courier New"


# ═══════════════════════════════════════════════════════════════════════════════
#  HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

def _field_char(type_):
    fc = OxmlElement("w:fldChar")
    fc.set(qn("w:fldCharType"), type_)
    return fc


def _instr_text(text):
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = text
    return it


def add_footer_page_num(doc):
    """Número de página centrado no rodapé."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run()
        r._r.append(_field_char("begin"))
        r._r.append(_instr_text("PAGE"))
        r._r.append(_field_char("end"))
        r.font.name = FONT
        r.font.size = Pt(10)


def insert_toc(doc):
    """Insere campo de Sumário automático (requer atualização no Word)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run()
    r._r.append(_field_char("begin"))
    r._r.append(_instr_text(' TOC \\o "1-3" \\h \\z \\u '))
    r._r.append(_field_char("separate"))
    # Texto de marcador (visível antes de atualizar)
    r2 = p.add_run("[ Clique com o botão direito neste campo e selecione "
                   "\"Atualizar Campo\" para gerar o Sumário ]")
    r2.font.name = FONT
    r2.font.size = Pt(10)
    r2.font.color.rgb = GRAY
    r._r.append(_field_char("end"))


def _run(para, text, bold=False, italic=False, size=12,
         color=None, mono=False):
    r = para.add_run(text)
    r.bold   = bold
    r.italic = italic
    r.font.name = MONO if mono else FONT
    r.font.size = Pt(10 if mono else size)
    if color:
        r.font.color.rgb = color
    return r


def body_para(doc, text, indent=True, bold=False, italic=False,
              size=12, color=None, mono=False,
              align=WD_ALIGN_PARAGRAPH.JUSTIFY,
              space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.style = "Normal"
    p.alignment = align
    p.paragraph_format.first_line_indent = Cm(1.25) if indent else Cm(0)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        _run(p, text, bold=bold, italic=italic, size=size,
             color=color, mono=mono)
    return p


def caption_para(doc, text):
    """Legenda de figura / tabela (tamanho 10, cinza, centralizada)."""
    p = doc.add_paragraph()
    p.style = "Normal"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(12)
    _run(p, text, size=10, color=GRAY, italic=True)
    return p


def note_para(doc, text):
    """Nota / fonte (recuada, tamanho 10)."""
    p = doc.add_paragraph()
    p.style = "Normal"
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent       = Cm(1.25)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(8)
    _run(p, text, size=10, color=GRAY)
    return p


def placeholder_para(doc, text):
    """Marcador de captura de tela (caixa cinza com borda)."""
    p = doc.add_paragraph()
    p.style = "Normal"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(10)
    _run(p, text, size=11, color=GRAY, italic=True)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "EFEFEF")
    pPr.append(shd)
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "4")
        border.set(qn("w:color"), "AAAAAA")
        pBdr.append(border)
    pPr.append(pBdr)
    return p


def render_image(doc, path):
    """Insere imagem centralizada; exibe marcador se arquivo ausente."""
    full_path = path if os.path.isabs(path) else os.path.join(BASE, path)
    if not os.path.exists(full_path):
        placeholder_para(doc, f"[IMAGEM NÃO ENCONTRADA: {path}]")
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.first_line_indent = Cm(0)
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after  = Pt(4)
    run = para.add_run()
    run.add_picture(full_path, width=Cm(14))


def code_para(doc, text):
    """Bloco de código / exemplo (monoespaçado, recuado)."""
    p = doc.add_paragraph()
    p.style = "Normal"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent       = Cm(1.25)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    _run(p, text, mono=True)
    return p


def heading(doc, text, level=1):
    """Heading 1 / 2 / 3 com formatação PUCRS."""
    p = doc.add_heading(text, level=level)
    p.paragraph_format.first_line_indent = Cm(0)
    sizes = {1: 16, 2: 13, 3: 12}
    for r in p.runs:
        r.font.name  = FONT
        r.font.size  = Pt(sizes.get(level, 12))
        r.font.bold  = True
        r.font.color.rgb = BLUE
    return p


# ═══════════════════════════════════════════════════════════════════════════════
#  CONFIGURAÇÃO DO DOCUMENTO
# ═══════════════════════════════════════════════════════════════════════════════

def setup_doc():
    doc = Document()

    # Margens ABNT/PUCRS
    s = doc.sections[0]
    s.page_height   = Cm(29.7)
    s.page_width    = Cm(21.0)
    s.top_margin    = Cm(3.0)
    s.bottom_margin = Cm(2.0)
    s.left_margin   = Cm(3.0)
    s.right_margin  = Cm(2.0)

    # Normal
    n = doc.styles["Normal"]
    n.font.name = FONT
    n.font.size = Pt(12)
    pf = n.paragraph_format
    pf.alignment          = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent  = Cm(1.25)
    pf.space_after        = Pt(0)
    pf.space_before       = Pt(0)
    pf.line_spacing       = Pt(18)   # 1.5 linhas

    # Heading 1
    h1 = doc.styles["Heading 1"]
    h1.font.name = FONT
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = BLUE
    h1.paragraph_format.space_before      = Pt(24)
    h1.paragraph_format.space_after       = Pt(12)
    h1.paragraph_format.first_line_indent = Cm(0)
    h1.paragraph_format.keep_with_next    = True
    h1.paragraph_format.page_break_before = True

    # Heading 2
    h2 = doc.styles["Heading 2"]
    h2.font.name = FONT
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = BLUE
    h2.paragraph_format.space_before      = Pt(18)
    h2.paragraph_format.space_after       = Pt(8)
    h2.paragraph_format.first_line_indent = Cm(0)
    h2.paragraph_format.keep_with_next    = True

    # Heading 3
    h3 = doc.styles["Heading 3"]
    h3.font.name = FONT
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = BLUE
    h3.paragraph_format.space_before      = Pt(12)
    h3.paragraph_format.space_after       = Pt(6)
    h3.paragraph_format.first_line_indent = Cm(0)

    add_footer_page_num(doc)
    return doc


# ═══════════════════════════════════════════════════════════════════════════════
#  CAPA
# ═══════════════════════════════════════════════════════════════════════════════

def add_cover(doc):
    def _cp(text, align=WD_ALIGN_PARAGRAPH.CENTER, bold=False,
            size=12, before=0, after=0, italic=False):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after  = Pt(after)
        p.paragraph_format.line_spacing = Pt(size * 1.5)
        _run(p, text, bold=bold, size=size, italic=italic)

    _cp("PONTIFÍCIA UNIVERSIDADE CATÓLICA DO RIO GRANDE DO SUL",
        bold=True, before=72)
    _cp("ESCOLA POLITÉCNICA", bold=True)
    _cp("BACHARELADO EM CIÊNCIA DA COMPUTAÇÃO", bold=True, after=0)

    _cp("BEDFLOW-ATLAS: Automação da geração de leitos empacotados\n"
        "com uso de software em código aberto",
        bold=True, size=16, before=80, after=0)

    _cp("Bernardo Klein Heitz", bold=True, size=14, before=40, after=0)

    # Descrição (alinhada à direita)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.line_spacing = Pt(18)
    _run(p, "Trabalho de Conclusão II apresentado como requisito\n"
            "parcial à obtenção do grau de Bacharel em Ciência da\n"
            "Computação na Pontifícia Universidade Católica do\n"
            "Rio Grande do Sul.", size=12)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.first_line_indent = Cm(0)
    p2.paragraph_format.space_before = Pt(30)
    p2.paragraph_format.space_after  = Pt(0)
    _run(p2, "Orientador:  Marco Aurélio Souza Mangan", size=12)

    _cp("Porto Alegre\n2026", bold=True, size=12, before=72)

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  PRÉ-TEXTUAIS (Siglas, Banco de Palavras, Sumário)
# ═══════════════════════════════════════════════════════════════════════════════

SIGLAS = [
    ("2D",       "Bidimensional"),
    ("3D",       "Tridimensional"),
    ("ANTLR",    "ANother Tool for Language Recognition"),
    ("API",      "Application Programming Interface"),
    ("CFD",      "Computational Fluid Dynamics"),
    ("CLI",      "Command-Line Interface"),
    ("CSV",      "Comma-Separated Values"),
    ("DB",       "Database"),
    ("DEM",      "Discrete Element Method"),
    ("DOE",      "Design of Experiments"),
    ("DOI",      "Digital Object Identifier"),
    ("DSL",      "Domain-Specific Language"),
    ("E2E",      "End to End"),
    ("FBX",      "Filmbox"),
    ("FOAM",     "Field Operation and Manipulation"),
    ("GCI",      "Grid Convergence Index"),
    ("GLB",      "GL Binary — formato binário do GLTF"),
    ("GLTF",     "GL Transmission Format"),
    ("GPU",      "Graphics Processing Unit"),
    ("GUI",      "Graphical User Interface"),
    ("HPC",      "High-Performance Computing"),
    ("HTTP",     "Hypertext Transfer Protocol"),
    ("IC",       "Iniciação Científica"),
    ("JSON",     "JavaScript Object Notation"),
    ("JWT",      "JSON Web Token"),
    ("LES",      "Large Eddy Simulation"),
    ("LOC",      "Lines of Code (linhas de código)"),
    ("LTS",      "Long-Term Support"),
    ("MIT",      "Licença de software de código aberto (Massachusetts Institute of Technology)"),
    ("MVP",      "Minimum Viable Product"),
    ("OBJ",      "Wavefront OBJ — formato de malha 3D"),
    ("OpenAPI",  "OpenAPI Specification"),
    ("OpenFOAM", "Open Field Operation and Manipulation"),
    ("ORM",      "Object-Relational Mapping"),
    ("PLY",      "Polygon File Format — formato de malha 3D"),
    ("PUCRS",    "Pontifícia Universidade Católica do Rio Grande do Sul"),
    ("RANS",     "Reynolds-Averaged Navier–Stokes"),
    ("RBAC",     "Role-Based Access Control"),
    ("RBD",      "Rigid Body Dynamics"),
    ("Re",       "Número de Reynolds"),
    ("REST",     "Representational State Transfer"),
    ("R3F",      "React Three Fiber"),
    ("S3",       "Simple Storage Service"),
    ("SHA",      "Secure Hash Algorithm"),
    ("SI",       "Sistema Internacional de Unidades"),
    ("SIC",      "Salão de Iniciação Científica"),
    ("SPA",      "Single Page Application"),
    ("SQL",      "Structured Query Language"),
    ("STL",      "Stereolithography / Standard Tessellation Language"),
    ("TCC",      "Trabalho de Conclusão de Curso"),
    ("UI",       "User Interface"),
    ("URL",      "Uniform Resource Locator"),
    ("VTK",      "Visualization Toolkit"),
    ("VTU",      "VTK Unstructured Grid"),
    ("YAML",     "YAML Ain't Markup Language"),
    ("WSL",      "Windows Subsystem for Linux"),
    ("Δp",       "Queda de pressão"),
    ("Δp/L",     "Queda de pressão por comprimento"),
    ("ε",        "Porosidade efetiva do leito"),
]

BANCO_PALAVRAS = [
    ("API",              "Interface de programação usada para comunicação entre sistemas por meio de endpoints HTTP."),
    ("Boolean Operation","Operação geométrica (diferença, união, interseção) entre duas meshes 3D para produzir uma geometria composta."),
    ("Bucket",           "Diretório lógico em storage de objetos (ex: MinIO, S3) onde arquivos são guardados."),
    ("Checksum",         "Valor calculado para verificar a integridade de um arquivo durante armazenamento ou transferência."),
    ("CSV",              "Formato texto com valores separados por vírgula, útil para métricas e tabelas."),
    ("Dashboard",        "Página web com gráficos e painéis que servem para acompanhar execuções, comparar resultados e baixar artefatos."),
    ("Endpoint",         "URL específica de uma API que expõe uma operação (criar job, listar execuções, obter arquivos)."),
    ("Fallback",         "Comportamento alternativo acionado automaticamente quando a opção principal não está disponível."),
    ("GCI",              "Índice para avaliar independência de malha e estimar erro numérico em estudos de refinamento."),
    ("Hash",             "Identificador derivado do conteúdo (ex: do params.json) usado para versionar e comparar variantes de execução."),
    ("Headless",         "Execução sem interface gráfica; no projeto, o Blender roda headless via scripts Python."),
    ("Job",              "Unidade de trabalho no pipeline, do .bed até os resultados consolidados."),
    ("Logs",             "Registros estruturados de eventos que servem para depuração, auditoria e rastreabilidade."),
    ("Manifold",         "Propriedade de superfícies sem buracos e com normais consistentes, adequada à malha CFD."),
    ("Mesh",             "Malha numérica que discretiza o domínio para resolver as equações de escoamento."),
    ("Modeling Profile", "Perfil de modelagem geométrica do sistema: Blender (com RBD) ou Python puro (stdlib)."),
    ("MVP",              "Primeira versão funcional mínima de um sistema, suficiente para demonstrar o fluxo fim a fim."),
    ("OpenFOAM",         "Conjunto de bibliotecas e solvers de CFD de código aberto usado neste trabalho."),
    ("Pipeline",         "Sequência organizada de etapas automatizadas que transforma a descrição do problema em resultados."),
    ("Queue",            "Fila de tarefas que organiza a execução de jobs, habilitando retries e paralelismo."),
    ("RANS",             "Família de modelos de turbulência baseada na média de Reynolds para fechar Navier–Stokes."),
    ("Seed",             "Valor inicial que controla processos estocásticos (ex: empacotamento), garantindo repetição exata."),
    ("Sidecar JSON",     "Arquivo JSON gerado junto à malha STL/OBJ com metadados completos da geração (modo, seed, porosidade etc.)."),
    ("Solver",           "Programa que resolve numericamente as equações do escoamento (ex: simpleFoam)."),
    ("STL",              "Formato de superfície triangulada usado na exportação de geometrias e importação em solvers CFD."),
    ("Stdlib",           "Biblioteca padrão de uma linguagem de programação, sem dependências externas (ex: math, struct, pathlib do Python)."),
    ("Thin Slice",       "Fatia fina do leito gerada pelo sistema de corte 3D paramétrico, resultando em geometria pseudo-2D."),
    ("URL",              "Localizador uniforme de recursos; usado para acessar endpoints e arquivos da API."),
    ("VTK",              "Ecossistema e formato para visualização científica, usado para exportar campos como pressão e velocidade."),
]


def add_siglas(doc):
    # Título
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.first_line_indent = Cm(0)
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after  = Pt(24)
    _run(p_title, "LISTA DE SIGLAS/SÍMBOLOS", bold=True, size=14)

    for sigla, definicao in SIGLAS:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.line_spacing = Pt(14)
        r1 = _run(p, sigla, italic=True, size=11)
        _run(p, "  ———  ", italic=True, size=11)
        _run(p, definicao, italic=True, size=11)

    doc.add_page_break()


def add_banco_palavras(doc):
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.first_line_indent = Cm(0)
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after  = Pt(24)
    _run(p_title, "BANCO DE PALAVRAS (TERMOS/SIGLAS EM INGLÊS)", bold=True, size=14)

    for termo, definicao in BANCO_PALAVRAS:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.line_spacing = Pt(14)
        _run(p, termo + "  —  ", italic=True, bold=True, size=11)
        _run(p, definicao, italic=True, size=11)

    doc.add_page_break()


def add_sumario(doc):
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.first_line_indent = Cm(0)
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after  = Pt(24)
    _run(p_title, "SUMÁRIO", bold=True, size=14)
    insert_toc(doc)
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  PARSER DO TXT
# ═══════════════════════════════════════════════════════════════════════════════

def parse_txt(path):
    """
    Retorna lista de dicionários representando blocos estruturados do TXT.
    Estratégia: máquina de estados linha a linha com detecção de padrões.
    """
    with open(path, encoding="utf-8") as f:
        lines = f.readlines()

    # ── Predicados ──────────────────────────────────────────────────────────

    def is_sep(ln):
        s = ln.strip()
        return s.startswith("===") and len(s) > 20

    def is_dash(ln):
        s = ln.strip()
        return s.startswith("---") and len(s) > 10

    def is_tbl_border(ln):
        return bool(re.match(r"^\s*\+[-+=+]+\+\s*$", ln))

    def is_tbl_row(ln):
        s = ln.strip()
        return s.startswith("|") and s.endswith("|") and s.count("|") >= 2

    def is_bullet(ln):
        s = ln.strip()
        return s.startswith("•") or s.startswith("·") or s.startswith("-  ")

    def is_enum(ln):
        s = ln.strip()
        return bool(re.match(
            r"^(OE|RF|RNF|CA|TF|TF\d|L|D)\d+[\.\s]", s))

    def is_code(ln):
        """Bloco de código: 4+ espaços e parece código (não prosa PT)."""
        if not ln.startswith("    "):
            return False
        s = ln.strip()
        if not s or is_tbl_border(ln) or is_tbl_row(ln) or is_bullet(ln) or is_enum(ln):
            return False
        # Começa com letra maiúscula e sem caracteres de código → prosa recuada
        if re.match(r"^[A-ZÁÀÃÉÊÍÓÕÚÇÜ][a-záàãéêíóõúçü]", s):
            return False
        return True

    def is_note(ln):
        s = ln.strip()
        return (s.startswith("*") or s.startswith("\"Net-new\"") or
                re.match(r"(?i)^(fonte:|nota (?:de |metodológica|:\s)|observa|diagrama\s+d)", s))

    def is_caption(ln):
        s = ln.strip()
        return bool(re.match(r"(?i)^figura\s+\d+|^figura\s+[A-Z]+\s+—", s))

    def is_image(ln):
        return ln.strip().startswith("[FIGURA:")

    def is_placeholder(ln):
        return bool(re.match(r"^\[\s*[WTR]\d+\s*\]", ln.strip()))

    # ── Ignorar pré-textuais; começar no cap. 1 ─────────────────────────────
    n = len(lines)
    start = 0
    for idx, ln in enumerate(lines):
        if re.match(r"^\s*1\s+INTRODUÇÃO", ln.strip()):
            # Buscar o === imediatamente antes
            j = idx - 1
            while j >= 0 and not is_sep(lines[j]):
                j -= 1
            start = j if j >= 0 else idx
            break

    blocks = []
    i = start

    while i < n:
        ln = lines[i]
        s  = ln.strip()

        # ── vazia ────────────────────────────────────────────────────────────
        if not s:
            i += 1
            continue

        # ── Separador === → capítulo ─────────────────────────────────────────
        if is_sep(ln):
            i += 1
            while i < n and not lines[i].strip():
                i += 1
            if i >= n:
                break
            title = lines[i].strip()
            i += 1
            # consumir === de fechamento
            while i < n and not lines[i].strip():
                i += 1
            if i < n and is_sep(lines[i]):
                i += 1

            skip_titles = {"DOCUMENTO INTEGRADO", "FIM DO DOCUMENTO",
                           "COMPLEMENTAÇÃO DAS REFERÊNCIAS"}
            if any(t in title.upper() for t in
                   {"LISTA DE SIGLAS", "GLOSSÁRIO", "BANCO DE PALAVRAS",
                    "COMPLEMENTAÇÃO", "FIM DO"}):
                # Pré-textuais / notas → ignorar
                continue

            ch = re.match(r"^(\d+)\s+(.+)", title)
            if ch or any(k in title.upper() for k in
                         ("REFERÊNCIA", "SEÇÃO ESPECIAL", "SIC 2026",
                          "ÍNDICE", "ARTEFATOS", "TABELAS CONSOLIDADAS")):
                blocks.append({"type": "chapter", "title": title, "level": 1})
            continue

        # ── Separador --- ─────────────────────────────────────────────────────
        if is_dash(ln):
            i += 1
            continue

        # ── Seção N.N (linha seguida de ---) ─────────────────────────────────
        sec = re.match(r"^(\d+\.\d+(?:\.\d+)?)\s+(.+)", s)
        if sec and i + 1 < n and is_dash(lines[i + 1]):
            # N.N → Heading 2 (subseção); N.N.N → Heading 3 (subsubseção)
            blocks.append({"type": "section", "title": s,
                            "level": min(s.count(".") + 1, 3)})
            i += 2
            continue

        # ── Tabela ASCII ──────────────────────────────────────────────────────
        if is_tbl_border(ln) or is_tbl_row(ln):
            tbl_lines = []
            while i < n and (is_tbl_border(lines[i]) or is_tbl_row(lines[i])
                              or (not lines[i].strip() and i + 1 < n
                                  and (is_tbl_border(lines[i+1])
                                       or is_tbl_row(lines[i+1])))):
                if lines[i].strip():
                    tbl_lines.append(lines[i])
                i += 1
            rows = _parse_ascii_table(tbl_lines)
            if rows:
                blocks.append({"type": "table", "rows": rows})
            continue

        # ── Lista com bullet (•) ──────────────────────────────────────────────
        if is_bullet(ln):
            items = []
            while i < n:
                sl = lines[i].strip()
                if not sl:
                    i += 1
                    if not items:
                        continue
                    break
                if is_bullet(lines[i]):
                    items.append(sl.lstrip("•·- ").strip())
                    i += 1
                elif items and lines[i].startswith("      "):
                    items[-1] += " " + sl
                    i += 1
                else:
                    break
            if items:
                blocks.append({"type": "list", "items": items})
            continue

        # ── Lista enumerada (OE1, RF01, …) ───────────────────────────────────
        if is_enum(ln):
            items = []
            while i < n:
                sl = lines[i].strip()
                if not sl:
                    i += 1
                    break
                if is_enum(lines[i]):
                    items.append(sl)
                    i += 1
                elif items and lines[i].startswith("      "):
                    # Continuação recuada
                    items[-1] += " " + sl
                    i += 1
                else:
                    break
            if items:
                blocks.append({"type": "list", "items": items, "ordered": True})
            continue

        # ── Bloco de código (4+ espaços, caracteres de código) ───────────────
        if is_code(ln):
            code_lines = []
            while i < n:
                cl = lines[i]
                if not cl.strip():
                    # Linha em branco: continua se próxima for código
                    if i + 1 < n and is_code(lines[i + 1]):
                        code_lines.append("")
                        i += 1
                        continue
                    else:
                        break
                if is_tbl_border(cl) or is_tbl_row(cl) or is_bullet(cl) or is_enum(cl):
                    break
                if is_sep(cl) or is_dash(cl):
                    break
                if is_code(cl) or cl.startswith("    "):
                    code_lines.append(cl.rstrip())
                    i += 1
                else:
                    break
            if code_lines:
                blocks.append({"type": "code", "lines": code_lines})
            continue

        # ── Imagem [FIGURA: filename] ─────────────────────────────────────────
        if is_image(ln):
            m = re.match(r"^\[FIGURA:\s*(.+?)\s*\]", s)
            path = m.group(1) if m else s
            blocks.append({"type": "image", "path": path})
            i += 1
            continue

        # ── Placeholder de captura [ W1 ] / [ T1 ] / [ R1 ] ─────────────────
        if is_placeholder(ln):
            blocks.append({"type": "placeholder", "text": s})
            i += 1
            continue

        # ── Legenda de figura ─────────────────────────────────────────────────
        if is_caption(ln):
            caption_text = s
            i += 1
            while i < n and lines[i].strip() and not is_sep(lines[i]):
                caption_text += " " + lines[i].strip()
                i += 1
            blocks.append({"type": "caption", "text": caption_text})
            continue

        # ── Nota / Fonte ──────────────────────────────────────────────────────
        if is_note(ln):
            note_text = s
            i += 1
            while i < n and lines[i].strip() and not is_sep(lines[i]) \
                  and not is_dash(lines[i]):
                nl = lines[i].strip()
                if is_caption(lines[i]) or is_enum(lines[i]) or is_bullet(lines[i]):
                    break
                note_text += " " + nl
                i += 1
            blocks.append({"type": "note", "text": note_text})
            continue

        # ── Parágrafo normal ──────────────────────────────────────────────────
        para_text = s
        i += 1
        while i < n:
            nl = lines[i].strip()
            if not nl:
                break
            if is_sep(lines[i]) or is_dash(lines[i]):
                break
            if re.match(r"^\d+\.\d+\s+", nl) and i + 1 < n and is_dash(lines[i+1]):
                break
            if is_tbl_border(lines[i]) or is_tbl_row(lines[i]):
                break
            if is_bullet(lines[i]) or is_enum(lines[i]):
                break
            if is_caption(lines[i]) or is_note(lines[i]):
                break
            para_text += " " + nl
            i += 1

        if para_text:
            blocks.append({"type": "para", "text": para_text})

    return blocks


def _parse_ascii_table(lines):
    """Converte linhas de tabela ASCII em lista de listas de strings."""
    rows = []
    current_cells = None

    for line in lines:
        s = line.strip()
        if not s:
            continue

        if re.match(r"^\+[-+=+]+\+$", s):
            if current_cells is not None:
                rows.append([c.strip() for c in current_cells])
                current_cells = None
        elif s.startswith("|") and s.endswith("|"):
            parts = s.split("|")
            cells = [p.strip() for p in parts[1:-1]]
            if current_cells is None:
                current_cells = cells[:]
            else:
                for idx, cell in enumerate(cells):
                    if idx < len(current_cells) and cell:
                        current_cells[idx] += " " + cell

    if current_cells is not None:
        rows.append([c.strip() for c in current_cells])

    return rows


# ═══════════════════════════════════════════════════════════════════════════════
#  RENDERIZADOR
# ═══════════════════════════════════════════════════════════════════════════════

def _render_table(doc, rows):
    """Renderiza tabela Word a partir de lista de listas."""
    if not rows:
        return

    n_cols = max(len(r) for r in rows)
    n_rows = len(rows)
    if n_cols == 0:
        return

    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Largura total disponível (16 cm = 21 - 3 - 2)
    col_width = Cm(16.0 / n_cols)

    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx]
        is_header = (r_idx == 0)
        for c_idx in range(n_cols):
            cell = row.cells[c_idx]
            cell.width = col_width
            text = row_data[c_idx] if c_idx < len(row_data) else ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)
            r = para.add_run(text)
            r.font.name = FONT
            r.font.size = Pt(9)
            r.bold = is_header
            if is_header:
                # Background cinza claro para cabeçalho
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"),   "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"),  "D9D9D9")
                tcPr.append(shd)

    # Espaço após tabela
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_after = Pt(8)


def render_blocks(doc, blocks):
    """Converte lista de blocos em conteúdo Word."""
    for blk in blocks:
        t = blk["type"]

        if t == "chapter":
            heading(doc, blk["title"], level=1)

        elif t == "section":
            heading(doc, blk["title"], level=blk.get("level", 2))

        elif t == "para":
            text = blk["text"]
            # Detectar se é citação direta (começa/termina com aspas ou itálico)
            is_quote = text.startswith('"') or text.startswith('“')
            body_para(doc, text, indent=True, italic=is_quote)

        elif t == "table":
            _render_table(doc, blk["rows"])

        elif t == "code":
            for line in blk["lines"]:
                code_para(doc, line)

        elif t == "list":
            ordered = blk.get("ordered", False)
            for item in blk["items"]:
                p = doc.add_paragraph()
                p.style = "Normal"
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Cm(-0.5)
                p.paragraph_format.left_indent       = Cm(1.75)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(3)
                prefix = "" if ordered else "• "
                _run(p, prefix + item, size=12)

        elif t == "note":
            note_para(doc, blk["text"])

        elif t == "caption":
            caption_para(doc, blk["text"])

        elif t == "image":
            render_image(doc, blk["path"])

        elif t == "placeholder":
            placeholder_para(doc, blk["text"])

        elif t == "page_break":
            doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  MAIN
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    print(f"Lendo: {TXT_IN}")
    if not os.path.exists(TXT_IN):
        print(f"ERRO: arquivo não encontrado: {TXT_IN}", file=sys.stderr)
        sys.exit(1)

    print("Configurando documento...")
    doc = setup_doc()

    print("Adicionando capa...")
    add_cover(doc)

    print("Adicionando Lista de Siglas...")
    add_siglas(doc)

    print("Adicionando Banco de Palavras...")
    add_banco_palavras(doc)

    print("Adicionando Sumário...")
    add_sumario(doc)

    print("Parseando conteúdo...")
    blocks = parse_txt(TXT_IN)
    print(f"  {len(blocks)} blocos identificados.")

    # Estatísticas
    type_counts = {}
    for b in blocks:
        type_counts[b["type"]] = type_counts.get(b["type"], 0) + 1
    for k, v in sorted(type_counts.items()):
        print(f"    {k}: {v}")

    print("Renderizando conteúdo...")
    render_blocks(doc, blocks)

    print(f"Salvando: {DOCX_OUT}")
    doc.save(DOCX_OUT)
    size_kb = os.path.getsize(DOCX_OUT) // 1024
    print(f"✓ Documento gerado com sucesso! ({size_kb} KB)")
    print(f"  Caminho: {DOCX_OUT}")


if __name__ == "__main__":
    main()
